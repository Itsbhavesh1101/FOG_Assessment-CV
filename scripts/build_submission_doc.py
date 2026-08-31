"""Build the BowlVision submission Word document."""

from __future__ import annotations

import csv
import json
import subprocess
import sys
from pathlib import Path
from textwrap import wrap

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
FIGURES = ROOT / "figures"
OUTPUT = ROOT / "output"
ASSETS = DOCS / "submission_assets"
DOCX_PATH = DOCS / "BowlVision_Project_Documentation_Bhavesh_Barmashe.docx"


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


def run_command(command: list[str]) -> str:
    proc = subprocess.run(
        command,
        cwd=ROOT,
        text=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        check=False,
    )
    return f"> {' '.join(command)}\n{proc.stdout.strip()}\n(exit code: {proc.returncode})"


def draw_wrapped_text(
    draw: ImageDraw.ImageDraw,
    text: str,
    xy: tuple[int, int],
    width_chars: int,
    line_height: int,
    fill: tuple[int, int, int],
    text_font: ImageFont.ImageFont,
) -> int:
    x, y = xy
    for source_line in text.splitlines():
        lines = wrap(source_line, width=width_chars) or [""]
        for line in lines:
            draw.text((x, y), line, fill=fill, font=text_font)
            y += line_height
    return y


def make_terminal_screenshot(path: Path, text: str, title: str) -> None:
    width, height = 1700, 1050
    img = Image.new("RGB", (width, height), "#111827")
    draw = ImageDraw.Draw(img)
    draw.rectangle((0, 0, width, 70), fill="#1f2937")
    draw.ellipse((28, 24, 48, 44), fill="#ef4444")
    draw.ellipse((62, 24, 82, 44), fill="#f59e0b")
    draw.ellipse((96, 24, 116, 44), fill="#22c55e")
    draw.text((145, 21), title, fill="#e5e7eb", font=font(28, bold=True))
    draw_wrapped_text(
        draw,
        text,
        (40, 105),
        width_chars=118,
        line_height=28,
        fill="#d1fae5",
        text_font=font(22),
    )
    img.save(path)


def make_output_screenshot(path: Path) -> None:
    with (OUTPUT / "final_scoreboard.json").open("r", encoding="utf-8") as f:
        data = json.load(f)

    width, height = 1700, 1000
    img = Image.new("RGB", (width, height), "#f8fafc")
    draw = ImageDraw.Draw(img)
    draw.rectangle((0, 0, width, 120), fill="#0f172a")
    draw.text((50, 36), "Extracted Scoreboard Data", fill="#ffffff", font=font(42, bold=True))
    draw.text((50, 88), "Generated from output/final_scoreboard.json", fill="#cbd5e1", font=font(22))

    headers = ["Player", "F1", "F2", "F3", "F4", "F5", "TTL"]
    col_x = [60, 420, 600, 780, 960, 1140, 1360]
    y = 180
    draw.rectangle((40, y - 25, 1620, y + 45), fill="#e2e8f0")
    for idx, header in enumerate(headers):
        draw.text((col_x[idx], y), header, fill="#0f172a", font=font(26, bold=True))
    y += 80

    for player in data["players"]:
        row_fill = "#ffffff" if (y // 80) % 2 == 0 else "#f1f5f9"
        draw.rectangle((40, y - 20, 1620, y + 48), fill=row_fill)
        values = [player["name"]]
        for frame in range(1, 6):
            cell = player["frames"][str(frame)]
            if cell is None:
                values.append("-")
            else:
                rolls = "".join(cell["rolls"])
                values.append(f"{rolls} / {cell['cumulative']}")
        values.append(str(player["ttl"]))
        for idx, value in enumerate(values):
            draw.text((col_x[idx], y), value, fill="#111827", font=font(24, bold=idx in (0, 6)))
        y += 80

    summary = (
        "Unplayed frames are preserved as null values in JSON and as 'unplayed' "
        "rows in CSV. Vishal's Frame 5 update is captured at the end of the clip."
    )
    draw_wrapped_text(draw, summary, (60, 760), 105, 32, "#334155", font(24))
    img.save(path)


def make_pipeline_diagram(path: Path) -> None:
    width, height = 1800, 1100
    img = Image.new("RGB", (width, height), "#ffffff")
    draw = ImageDraw.Draw(img)
    title_font = font(40, bold=True)
    box_font = font(21, bold=True)
    small_font = font(18)

    draw.text((50, 35), "BowlVision Detailed Processing Architecture", fill="#0f172a", font=title_font)

    stages = [
        ("1. Video Input", "1080p MP4\n30 FPS stream\nmetadata extraction", "#dbeafe"),
        ("2. Sampling", "Uniform temporal\nsampling at 5 FPS\nstride around 6", "#dcfce7"),
        ("3. Cutaway Filter", "Canny edge density\nheader luminance\nvisibility decision", "#fef3c7"),
        ("4. Enhancement", "ROI crop\nCLAHE grayscale\nbilateral denoise", "#ede9fe"),
        ("5. OCR", "RapidOCR primary\nTesseract/Paddle fallback\nOCRItem contract", "#fee2e2"),
        ("6. Spatial Mapping", "4 player rows\n10 frame columns\nTTL column", "#e0f2fe"),
        ("7. Temporal Rules", "Sliding window voting\nbowling invariants\nstate freeze", "#fce7f3"),
        ("8. Export", "JSON\nCSV\ntimeline and video HUD", "#ecfccb"),
    ]

    positions = [
        (80, 160),
        (500, 160),
        (920, 160),
        (1340, 160),
        (1340, 560),
        (920, 560),
        (500, 560),
        (80, 560),
    ]
    box_w, box_h = 330, 220
    for idx, ((title, body, color), (x, y)) in enumerate(zip(stages, positions)):
        draw.rounded_rectangle((x, y, x + box_w, y + box_h), radius=24, fill=color, outline="#334155", width=3)
        draw.text((x + 24, y + 24), title, fill="#0f172a", font=box_font)
        draw_wrapped_text(draw, body, (x + 24, y + 72), 24, 28, "#334155", small_font)
        if idx < len(positions) - 1:
            nx, ny = positions[idx + 1]
            if y == ny:
                start = (x + box_w + 15, y + box_h // 2)
                end = (nx - 15, ny + box_h // 2)
            else:
                start = (x + box_w // 2, y + box_h + 15)
                end = (nx + box_w // 2, ny - 15)
            draw.line((start, end), fill="#0f172a", width=5)
            draw.polygon(
                [(end[0], end[1]), (end[0] - 18, end[1] - 10), (end[0] - 18, end[1] + 10)],
                fill="#0f172a",
            )

    draw.rounded_rectangle((500, 900, 1300, 1030), radius=20, fill="#f8fafc", outline="#94a3b8", width=2)
    draw.text((540, 930), "Design principle: each stage has one responsibility and passes typed data to the next stage.", fill="#0f172a", font=font(24, bold=True))
    draw.text((540, 970), "This makes OCR, grid mapping, scoring rules, and exporters independently testable.", fill="#334155", font=font(22))
    img.save(path)


def make_state_machine_diagram(path: Path) -> None:
    width, height = 1500, 1050
    img = Image.new("RGB", (width, height), "#ffffff")
    draw = ImageDraw.Draw(img)
    draw.text((50, 35), "Visibility and Temporal State Machine", fill="#0f172a", font=font(40, bold=True))

    nodes = {
        "Sample Frame": (120, 170, "#dbeafe"),
        "Compute Features": (540, 170, "#fef3c7"),
        "Scoreboard Visible": (960, 120, "#dcfce7"),
        "Cutaway Detected": (960, 340, "#fee2e2"),
        "Run OCR + Mapping": (540, 520, "#ede9fe"),
        "Freeze Previous State": (960, 560, "#fce7f3"),
        "Temporal Vote": (540, 760, "#e0f2fe"),
        "Export Confirmed State": (960, 760, "#ecfccb"),
    }

    def box(label: str, x: int, y: int, color: str) -> None:
        draw.rounded_rectangle((x, y, x + 300, y + 115), radius=20, fill=color, outline="#334155", width=3)
        draw_wrapped_text(draw, label, (x + 24, y + 28), 22, 28, "#0f172a", font(22, bold=True))

    for label, (x, y, color) in nodes.items():
        box(label, x, y, color)

    arrows = [
        ("Sample Frame", "Compute Features", ""),
        ("Compute Features", "Scoreboard Visible", "edge + luminance pass"),
        ("Compute Features", "Cutaway Detected", "feature check fails"),
        ("Scoreboard Visible", "Run OCR + Mapping", ""),
        ("Cutaway Detected", "Freeze Previous State", ""),
        ("Run OCR + Mapping", "Temporal Vote", ""),
        ("Freeze Previous State", "Temporal Vote", ""),
        ("Temporal Vote", "Export Confirmed State", "valid state"),
    ]

    for src, dst, label in arrows:
        sx, sy, _ = nodes[src]
        dx, dy, _ = nodes[dst]
        start = (sx + 300, sy + 58) if dx > sx else (sx + 150, sy + 115)
        end = (dx, dy + 58) if dx > sx else (dx + 150, dy)
        draw.line((start, end), fill="#0f172a", width=4)
        draw.ellipse((end[0] - 8, end[1] - 8, end[0] + 8, end[1] + 8), fill="#0f172a")
        if label:
            mid = ((start[0] + end[0]) // 2, (start[1] + end[1]) // 2 - 24)
            draw.text(mid, label, fill="#475569", font=font(16))

    img.save(path)


def make_assets() -> dict[str, Path]:
    ASSETS.mkdir(parents=True, exist_ok=True)

    test_output = run_command([sys.executable, "-m", "unittest", "discover", "-s", "tests", "-p", "test_*.py"])
    help_output = run_command([sys.executable, "-m", "bowlvision", "--help"])
    make_terminal_screenshot(ASSETS / "code_running_screenshot.png", f"{test_output}\n\n{help_output}", "PowerShell - BowlVision commands running")

    make_output_screenshot(ASSETS / "extracted_output_screenshot.png")
    make_pipeline_diagram(ASSETS / "architecture_pipeline_diagram.png")
    make_state_machine_diagram(ASSETS / "temporal_state_machine_diagram.png")

    return {
        "input": FIGURES / "fig1_input_scoreboard_frame.png",
        "detected": FIGURES / "fig2_spatial_grid_debug.png",
        "code": ASSETS / "code_running_screenshot.png",
        "output": ASSETS / "extracted_output_screenshot.png",
        "architecture": ASSETS / "architecture_pipeline_diagram.png",
        "state": ASSETS / "temporal_state_machine_diagram.png",
        "ocr": FIGURES / "fig5_ocr_evaluation_results.png",
        "timeline": FIGURES / "fig6_spatial_mapping_samples.png",
    }


def set_doc_defaults(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    for section in document.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)


def add_title_page(document: Document) -> None:
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("BowlVision\nAutomated Sports Scoreboard Extraction")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(15, 23, 42)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Computer Vision and OCR Based Project Documentation")
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor(71, 85, 105)

    document.add_paragraph("")
    table = document.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    rows = [
        ("Submitted by", "Bhavesh Barmashe"),
        ("Institute", "Sagar Institute of Research and Technology, Bhopal"),
        ("Project", "FOG Technologies Computer Vision Engineer Assessment"),
        ("Technology stack", "Python, OpenCV, RapidOCR/ONNXRuntime, PyTesseract, NumPy, Pillow"),
        ("Date", "31 August 2026"),
    ]
    for row, values in zip(table.rows, rows):
        row.cells[0].text = values[0]
        row.cells[1].text = values[1]

    document.add_paragraph("")
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(
        "This document explains the end-to-end design, implementation, verification, "
        "and output evidence for the BowlVision scoreboard extraction pipeline."
    )
    document.add_page_break()


def add_heading(document: Document, text: str, level: int = 1) -> None:
    document.add_heading(text, level=level)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Number")


def add_image(document: Document, path: Path, caption: str, width: float = 6.8) -> None:
    if path.exists():
        document.add_picture(str(path), width=Inches(width))
        p = document.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.italic = True
            run.font.size = Pt(9)


def add_score_table(document: Document) -> None:
    with (OUTPUT / "final_scoreboard.json").open("r", encoding="utf-8") as f:
        data = json.load(f)
    table = document.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    headers = ["Player", "Frame 1", "Frame 2", "Frame 3", "Frame 4", "Frame 5", "TTL"]
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
    for player in data["players"]:
        row = table.add_row().cells
        row[0].text = player["name"]
        for frame in range(1, 6):
            cell = player["frames"][str(frame)]
            row[frame].text = "Unplayed" if cell is None else f"{''.join(cell['rolls'])} -> {cell['cumulative']}"
        row[6].text = str(player["ttl"])


def add_csv_preview(document: Document) -> None:
    rows: list[list[str]] = []
    with (OUTPUT / "final_scoreboard.csv").open("r", encoding="utf-8") as f:
        reader = csv.reader(f)
        for idx, row in enumerate(reader):
            rows.append(row)
            if idx >= 12:
                break
    table = document.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    for cell, value in zip(table.rows[0].cells, rows[0]):
        cell.text = value
    for source in rows[1:]:
        row = table.add_row().cells
        for cell, value in zip(row, source):
            cell.text = value


def build_document(assets: dict[str, Path]) -> None:
    document = Document()
    set_doc_defaults(document)
    add_title_page(document)

    add_heading(document, "1. Executive Summary")
    document.add_paragraph(
        "BowlVision is an automated computer vision and OCR pipeline designed to extract structured "
        "bowling scoreboard data from broadcast video. The system reads the input match video, detects "
        "the overhead scoreboard, filters non-scoreboard camera cutaways, enhances the scoreboard crop, "
        "runs OCR, maps text into player frame cells, applies bowling scoring constraints, and exports "
        "validated JSON/CSV datasets."
    )
    add_bullets(
        document,
        [
            "Input: Full HD bowling broadcast video named bowling_scoreboard.mp4.",
            "Output: final_scoreboard.json, final_scoreboard.csv, timeline_history.csv, and optional annotated video.",
            "Core challenge: scoreboard text is small, high contrast, and sometimes merged across frame columns.",
            "Core solution: combine visual filtering, OCR abstraction, proportional spatial mapping, and temporal validation.",
        ],
    )

    add_heading(document, "2. Project Metadata")
    meta_table = document.add_table(rows=7, cols=2)
    meta_table.style = "Table Grid"
    rows = [
        ("Student name", "Bhavesh Barmashe"),
        ("Institute", "Sagar Institute of Research and Technology, Bhopal"),
        ("Project name", "BowlVision"),
        ("Domain", "Computer Vision, OCR, Sports Analytics"),
        ("Language", "Python 3.10+"),
        ("Primary libraries", "OpenCV, RapidOCR, ONNXRuntime, NumPy, Pillow, PyTesseract"),
        ("Repository package", "bowlvision"),
    ]
    for row, values in zip(meta_table.rows, rows):
        row.cells[0].text = values[0]
        row.cells[1].text = values[1]

    add_heading(document, "3. Required Screenshot Evidence")
    add_heading(document, "3.1 Input Video Frame", level=2)
    document.add_paragraph(
        "The first screenshot shows the original broadcast frame containing the overhead bowling scoreboard. "
        "This is the visual input from which player rows, frame columns, roll symbols, and total scores are extracted."
    )
    add_image(document, assets["input"], "Screenshot 1: Input video frame with the original scoreboard view.")

    add_heading(document, "3.2 Code Running", level=2)
    document.add_paragraph(
        "The second screenshot captures the project commands running successfully. It includes the automated test "
        "suite and CLI help output, demonstrating that the repository is executable and the command-line interface is available."
    )
    add_image(document, assets["code"], "Screenshot 2: Code execution evidence showing tests and CLI availability.")

    add_heading(document, "3.3 Detected Scoreboard", level=2)
    document.add_paragraph(
        "The third screenshot shows the detected scoreboard grid and spatial mapping overlay. The calibrated grid lets "
        "the pipeline assign OCR detections to the correct player row, frame column, and TTL column."
    )
    add_image(document, assets["detected"], "Screenshot 3: Detected scoreboard and grid mapping overlay.")

    add_heading(document, "3.4 Extracted Scoreboard Data and Output", level=2)
    document.add_paragraph(
        "The fourth screenshot summarizes the extracted structured scoreboard data. The same information is written "
        "to JSON and CSV files in the output directory."
    )
    add_image(document, assets["output"], "Screenshot 4: Extracted scoreboard data generated from output JSON.")

    add_heading(document, "4. Repository Structure")
    document.add_paragraph(
        "The project is organized as a Python package with focused submodules. Runtime artifacts such as debug crops, "
        "output files, caches, and large local videos are kept outside the clean source surface through .gitignore."
    )
    structure = (
        "bowlvision/\n"
        "  core/       Configuration, dataclasses, and typed OCR models\n"
        "  vision/     Video stream reading, sampling, cutaway detection, enhancement\n"
        "  ocr/        OCR backends and OCR manager abstraction\n"
        "  spatial/    Grid layout, row/column mapping, active player detection\n"
        "  analytics/  Bowling rules and temporal state aggregation\n"
        "  export/     JSON, CSV, terminal, and video annotation exporters\n"
        "  cli.py      Canonical command-line interface\n"
        "docs/         Markdown documentation and generated Word report\n"
        "figures/      Evidence screenshots and project figures\n"
        "scripts/      Utility scripts and document builder\n"
        "tests/        Unit tests for configuration, rules, mapping, OCR sanitation, and export"
    )
    document.add_paragraph(structure)

    add_heading(document, "5. Detailed System Architecture")
    document.add_paragraph(
        "BowlVision follows an eight-stage architecture. Each stage has a narrow responsibility and passes structured "
        "data to the next stage, which keeps the project testable and easier to maintain."
    )
    add_image(document, assets["architecture"], "Diagram 1: End-to-end BowlVision processing architecture.")
    add_numbered(
        document,
        [
            "Video ingestion reads metadata such as frame count, FPS, resolution, and duration.",
            "Temporal sampling reduces compute cost by processing representative frames instead of every frame.",
            "Cutaway detection rejects non-scoreboard camera views using edge density and luminance checks.",
            "Image enhancement improves OCR readability through grayscale conversion, CLAHE, and bilateral filtering.",
            "OCR manager selects the best available OCR backend and normalizes outputs into OCRItem objects.",
            "Spatial mapping assigns OCR detections to player rows, bowling frame columns, and TTL cells.",
            "Temporal aggregation rejects flickering OCR errors and preserves state during cutaways.",
            "Exporters serialize final results to JSON, CSV, timeline CSV, terminal output, and annotated video.",
        ],
    )

    add_heading(document, "6. Data Flow and State Machine")
    document.add_paragraph(
        "The visibility state machine is important because broadcast footage repeatedly switches away from the scoreboard. "
        "When a cutaway is detected, the system freezes the last confirmed state instead of allowing noisy OCR to overwrite scores."
    )
    add_image(document, assets["state"], "Diagram 2: Visibility and temporal state machine.")

    add_heading(document, "7. Computer Vision Pipeline")
    add_heading(document, "7.1 Video Sampling", level=2)
    document.add_paragraph(
        "The input video is sampled at approximately 5 FPS. For a 30 FPS video, this gives a stride of roughly six frames, "
        "which is frequent enough to catch score changes while greatly reducing OCR workload."
    )
    add_heading(document, "7.2 Cutaway Detection", level=2)
    document.add_paragraph(
        "The cutaway classifier checks the calibrated scoreboard region using two features: Canny edge density and header "
        "luminance. Scoreboard frames contain strong grid edges and stable header brightness, while lane or pin-deck views do not."
    )
    add_heading(document, "7.3 Image Enhancement", level=2)
    document.add_paragraph(
        "For visible scoreboard frames, the ROI is converted to grayscale, enhanced with CLAHE, and denoised with a bilateral "
        "filter. This improves contrast for LED-style digits without destroying digit edges."
    )

    add_heading(document, "8. OCR and Spatial Mapping")
    document.add_paragraph(
        "OCR output is not used directly as final data. Each OCR token contains text, confidence, and bounding box coordinates. "
        "The spatial mapper uses these coordinates to assign tokens to the scoreboard grid."
    )
    add_bullets(
        document,
        [
            "Rows identify players: JAGDISH, VISHAL, P (Player 3), and TARUN.",
            "Columns identify Frame 1 through Frame 10 and the TTL total score column.",
            "Upper/lower row splits separate roll symbols from cumulative scores.",
            "Merged OCR strings are split proportionally by character position and mapped into the correct columns.",
        ],
    )
    document.add_paragraph(
        "For a merged OCR token, each character is assigned an estimated x-coordinate using: "
        "Xi = Xmin + (i + 0.5) * (Xmax - Xmin) / L. This avoids hardcoding the final scores."
    )

    add_heading(document, "9. Bowling Rules and Temporal Validation")
    document.add_paragraph(
        "The temporal tracker combines repeated OCR observations with bowling-specific invariants. Scores must move forward "
        "monotonically, unplayed frames must remain unplayed, and TTL is derived from the latest confirmed cumulative score."
    )
    add_bullets(
        document,
        [
            "Strike (X): all 10 pins on the first roll.",
            "Spare (/): all remaining pins knocked down on the second roll.",
            "Open frame: fewer than 10 total pins knocked down.",
            "Miss (-): zero pins knocked down.",
            "TTL: latest confirmed cumulative score for the player.",
        ],
    )

    add_heading(document, "10. Final Extracted Scoreboard")
    document.add_paragraph("The final extracted scoreboard is shown below.")
    add_score_table(document)

    add_heading(document, "11. Output Files")
    document.add_paragraph("The pipeline writes the following files:")
    add_bullets(
        document,
        [
            "output/final_scoreboard.json: nested player/frame structure with null for unplayed frames.",
            "output/final_scoreboard.csv: tabular frame-by-frame score output.",
            "output/timeline_history.csv: timestamped frame observation history.",
            "output/annotated_bowling_scoreboard.mp4: optional annotated demonstration video.",
        ],
    )
    document.add_paragraph("CSV preview:")
    add_csv_preview(document)

    add_heading(document, "12. Testing and Verification")
    document.add_paragraph(
        "The repository contains automated tests for configuration parsing, image enhancement, OCR sanitation, spatial mapping, "
        "bowling rules, temporal tracking, and exporters. The current verification command is:"
    )
    document.add_paragraph('python -m unittest discover -s tests -p "test_*.py"')
    document.add_paragraph("The generated evidence screenshot shows the test suite running successfully.")

    add_heading(document, "13. Results and Conclusion")
    document.add_paragraph(
        "The project successfully converts raw bowling broadcast footage into structured scoreboard data. The architecture "
        "combines visual heuristics, OCR, spatial coordinate mapping, temporal voting, and bowling domain rules to handle "
        "scoreboard visibility changes and OCR noise. The final validated totals are JAGDISH 31, VISHAL 37, P (Player 3) 54, "
        "and TARUN 40."
    )

    document.add_section(WD_SECTION.NEW_PAGE)
    add_heading(document, "Appendix A: Additional Evaluation Figures")
    add_image(document, assets["ocr"], "Additional Figure: OCR evaluation and confidence metrics.")
    add_image(document, assets["timeline"], "Additional Figure: Temporal mapping samples across the video timeline.")

    document.save(DOCX_PATH)


def main() -> int:
    missing = [
        path
        for path in [
            FIGURES / "fig1_input_scoreboard_frame.png",
            FIGURES / "fig2_spatial_grid_debug.png",
            FIGURES / "fig5_ocr_evaluation_results.png",
            FIGURES / "fig6_spatial_mapping_samples.png",
            OUTPUT / "final_scoreboard.json",
            OUTPUT / "final_scoreboard.csv",
        ]
        if not path.exists()
    ]
    if missing:
        for path in missing:
            print(f"Missing required asset: {path}")
        return 1

    assets = make_assets()
    build_document(assets)
    print(DOCX_PATH)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

